#!/usr/bin/env python3
"""
Enhanced FastAPI Backend with RAG for Document Question Answering
Uses FAISS for vector search, OpenRouter for generation, and Supabase for storage
Production-ready with intelligent chunking and semantic retrieval - NO LOCAL STORAGE
"""

from fastapi import FastAPI, HTTPException, Depends, status, Request, UploadFile, File
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel, field_validator
from typing import List, Optional, Dict, Any, Tuple
from vector_cache import get_vector_cache
import httpx
import asyncio
import logging
import os
import sys
from dotenv import load_dotenv
import time
import hashlib
import json
import uuid
from pathlib import Path
import pickle
import io
from io import BytesIO
import signal
import tempfile
import mimetypes

# Document processing imports
import PyPDF2
import docx
import email

# RAG and embedding imports
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
import nltk
from nltk.tokenize import sent_tokenize
import re
from urllib.parse import urlparse, parse_qs

# OpenRouter LLM integration
from openrouter_service import OpenRouterService, AVAILABLE_MODELS
from session_manager import get_session_manager, TemporarySession

# Supabase storage utilities
from supabase_utils import (
    upload_file_to_supabase,
    download_file_from_supabase,
    download_document_content,
    get_public_url,
    delete_file_from_supabase,
    get_supabase_manager,
    list_supabase_files,
    test_supabase_upload_standalone
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - [%(filename)s:%(lineno)d] - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)

# Download required NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    try:
        nltk.download('punkt', quiet=True)
    except:
        try:
            nltk.download('punkt_tab', quiet=True)
        except:
            logger.warning("Could not download NLTK punkt tokenizer")
            pass


# Graceful shutdown handler for Cloud Run
class GracefulKiller:
    """Handle graceful shutdown for Cloud Run"""
    kill_now = False

    def __init__(self):
        signal.signal(signal.SIGINT, self._exit_gracefully)
        signal.signal(signal.SIGTERM, self._exit_gracefully)

    def _exit_gracefully(self, signum, frame):
        logger.info(f"Received signal {signum}, shutting down gracefully...")
        self.kill_now = True


killer = GracefulKiller()

# Load environment variables
load_dotenv()

# Initialize FastAPI app
app = FastAPI(
    title="RAG-Powered Document QA API with OpenRouter",
    description="Advanced Document Question Answering with Retrieval-Augmented Generation using OpenRouter",
    version="3.1.0",
    docs_url="/api/v1/docs",
    redoc_url="/api/v1/redoc"
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security
security = HTTPBearer()
VALID_TOKEN = os.getenv("VALID_TOKEN")

# Configuration
EMBEDDING_MODEL_NAME = "all-MiniLM-L6-v2"
CHUNK_SIZE = 1500
CHUNK_OVERLAP = 200
TOP_K_RETRIEVAL = 8
MAX_CONTEXT_LENGTH = 10000

# Global service instances (initialized in startup_event)
vector_store = None
llm_service = None

# Log environment variables
logger.info("🔧 ENVIRONMENT VARIABLES DEBUG:")
logger.info(f"   SUPABASE_URL: {os.getenv('SUPABASE_URL')}")
logger.info(
    f"   SUPABASE_KEY: {'*' * (len(os.getenv('SUPABASE_KEY', '')) - 8) + os.getenv('SUPABASE_KEY', '')[-8:] if os.getenv('SUPABASE_KEY') else 'NOT_SET'}")
logger.info(f"   SUPABASE_BUCKET: {os.getenv('SUPABASE_BUCKET', 'documents')}")
logger.info(f"   OPENROUTER_API_KEY: {'SET' if os.getenv('OPENROUTER_API_KEY') else 'NOT_SET'}")
logger.info(f"   OPENROUTER_MODEL: {os.getenv('OPENROUTER_MODEL', 'meta-llama/llama-3.1-8b-instruct:free')}")
logger.info(f"   VALID_TOKEN: {'SET' if VALID_TOKEN else 'NOT_SET'}")

# Validate environment variables
if not VALID_TOKEN:
    raise ValueError("VALID_TOKEN environment variable is required")
if not os.getenv("SUPABASE_URL"):
    raise ValueError("SUPABASE_URL environment variable is required")
if not os.getenv("SUPABASE_KEY"):
    raise ValueError("SUPABASE_KEY environment variable is required")
if not os.getenv("OPENROUTER_API_KEY"):
    raise ValueError("OPENROUTER_API_KEY environment variable is required. Get your key at https://openrouter.ai/keys")

class TemporarySessionCreate(BaseModel):
    ttl_minutes: int = 60

    @field_validator('ttl_minutes')
    @classmethod
    def validate_ttl(cls, v):
        if v < 5 or v > 240:
            raise ValueError("TTL must be between 5 and 240 minutes")
        return v


class TemporaryUploadRequest(BaseModel):
    session_id: str

    @field_validator('session_id')
    @classmethod
    def validate_session_id(cls, v):
        if not v or not v.strip():
            raise ValueError("Session ID cannot be empty")
        return v.strip()


class TemporaryDocumentQARequest(BaseModel):
    session_id: str
    questions: List[str]
    document_id: Optional[str] = None

    @field_validator('questions')
    @classmethod
    def validate_questions(cls, v):
        if not v or len(v) == 0:
            raise ValueError("At least one question is required")
        if len(v) > 10:
            raise ValueError("Maximum 10 questions allowed per request")
        return v

    @field_validator('session_id')
    @classmethod
    def validate_session_id(cls, v):
        if not v or not v.strip():
            raise ValueError("Session ID cannot be empty")
        return v.strip()

# Request/Response Models
class DocumentQARequest(BaseModel):
    documents: Optional[str] = None
    questions: List[str]
    document_id: Optional[str] = None

    @field_validator('questions')
    @classmethod
    def validate_questions(cls, v):
        if not v or len(v) == 0:
            raise ValueError("At least one question is required")
        if len(v) > 10:
            raise ValueError("Maximum 10 questions allowed per request")
        return v


class GlobalQueryRequest(BaseModel):
    query: str
    top_k: int = 10
    max_docs: int = 5
    document_ids: Optional[List[str]] = None

    @field_validator('query')
    @classmethod
    def validate_query(cls, v):
        if not v or not v.strip():
            raise ValueError("Query cannot be empty")
        if len(v) > 500:
            raise ValueError("Query too long (max 500 characters)")
        return v.strip()

    @field_validator('document_ids')
    @classmethod
    def validate_document_ids(cls, v):
        if v is not None:
            if len(v) == 0:
                raise ValueError("document_ids cannot be empty if provided")
            if len(v) > 20:
                raise ValueError("Maximum 20 documents allowed per query")
        return v


class DocumentUploadResponse(BaseModel):
    document_id: str
    filename: str
    status: str
    chunks_created: int
    message: str
    supabase_path: str
    public_url: Optional[str] = None


class DocumentQAResponse(BaseModel):
    answers: List[Dict[str, Any]]
    document_id: str
    retrieval_info: Optional[Dict[str, Any]] = None


class ErrorResponse(BaseModel):
    error: str
    details: Optional[str] = None


class DocumentIngestRequest(BaseModel):
    source: str
    source_type: Optional[str] = "auto"
    folder_name: Optional[str] = None

    @field_validator('source')
    @classmethod
    def validate_source(cls, v):
        if not v or not v.strip():
            raise ValueError("Source cannot be empty")
        return v.strip()


# Authentication dependency
async def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    if credentials.credentials != VALID_TOKEN:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid authentication token",
            headers={"WWW-Authenticate": "Bearer"},
        )
    return credentials.credentials


class DocumentProcessor:
    """Enhanced document processor with intelligent chunking and Supabase integration"""

    @staticmethod
    def extract_text_from_pdf(content: bytes) -> str:
        """Extract text from PDF with better error handling"""
        try:
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""

            for page_num in range(len(pdf_reader.pages)):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    if page_text.strip():
                        text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                except Exception as e:
                    logger.warning(f"Error processing page {page_num + 1}: {e}")
                    continue

            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Failed to process PDF: {str(e)}"
            )

    @staticmethod
    def extract_text_from_docx(content: bytes) -> str:
        """Extract text from DOCX with enhanced processing"""
        try:
            docx_file = BytesIO(content)
            doc = docx.Document(docx_file)
            text = ""

            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"

            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {e}")
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Failed to process DOCX: {str(e)}"
            )

    @staticmethod
    def extract_text_from_email(content: bytes) -> str:
        """Extract text from email content"""
        try:
            email_str = content.decode('utf-8', errors='ignore')
            msg = email.message_from_string(email_str)

            text = ""

            # Extract headers
            for header in ['From', 'To', 'Subject', 'Date']:
                if msg.get(header):
                    text += f"{header}: {msg.get(header)}\n"
            text += "\n"

            # Extract body
            if msg.is_multipart():
                for part in msg.walk():
                    if part.get_content_type() == "text/plain":
                        payload = part.get_payload(decode=True)
                        if payload:
                            text += payload.decode('utf-8', errors='ignore') + "\n"
            else:
                payload = msg.get_payload(decode=True)
                if payload:
                    text += payload.decode('utf-8', errors='ignore')

            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting email text: {e}")
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail=f"Failed to process email: {str(e)}"
            )

    @classmethod
    def intelligent_chunking(cls, text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[
        Dict[str, Any]]:
        """Intelligent text chunking that preserves semantic meaning"""
        # Clean the text
        text = cls._clean_text(text)

        # Split into sentences
        sentences = sent_tokenize(text)

        chunks = []
        current_chunk = ""
        current_chunk_sentences = []

        for i, sentence in enumerate(sentences):
            # Check if adding this sentence would exceed chunk size
            potential_chunk = current_chunk + " " + sentence if current_chunk else sentence

            if len(potential_chunk) <= chunk_size:
                current_chunk = potential_chunk
                current_chunk_sentences.append(sentence)
            else:
                # Save current chunk if it has content
                if current_chunk:
                    chunks.append({
                        "text": current_chunk.strip(),
                        "chunk_id": len(chunks),
                        "sentence_count": len(current_chunk_sentences),
                        "char_count": len(current_chunk)
                    })

                # Start new chunk with overlap
                if overlap > 0 and current_chunk_sentences:
                    # Calculate how many sentences to include for overlap
                    overlap_sentences = []
                    overlap_chars = 0

                    for sent in reversed(current_chunk_sentences):
                        if overlap_chars + len(sent) <= overlap:
                            overlap_sentences.insert(0, sent)
                            overlap_chars += len(sent)
                        else:
                            break

                    current_chunk = " ".join(overlap_sentences + [sentence])
                    current_chunk_sentences = overlap_sentences + [sentence]
                else:
                    current_chunk = sentence
                    current_chunk_sentences = [sentence]

        # Add the last chunk
        if current_chunk:
            chunks.append({
                "text": current_chunk.strip(),
                "chunk_id": len(chunks),
                "sentence_count": len(current_chunk_sentences),
                "char_count": len(current_chunk)
            })

        logger.info(f"Created {len(chunks)} intelligent chunks")
        return chunks

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""

        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[\f\r]+', '\n', text)
        text = re.sub(r'["""]', '"', text)
        text = re.sub(r'[\u2018\u2019]', "'", text)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)

        return text.strip()

    @classmethod
    async def process_document_from_source(cls, source: str, source_type: str = "auto") -> Tuple[str, str, str, str]:
        """Process document from URL or Supabase storage"""
        doc_id = hashlib.md5(source.encode()).hexdigest()

        if source_type == "auto":
            if GoogleDriveProcessor.is_google_drive_link(source):
                source_type = "google_drive"
            elif source.startswith(('http://', 'https://')):
                source_type = "url"
            else:
                source_type = "supabase_path"

        if source_type == "google_drive":
            original_filename, sanitized_filename = await cls._extract_filename_from_url(source)
            direct_url = GoogleDriveProcessor.convert_to_direct_download(source)
            source = direct_url
        else:
            original_filename, sanitized_filename = await cls._extract_filename_from_url(source)

        content, _ = await download_document_content(source)
        logger.info(f"Processing document: {original_filename} ({len(content)} bytes)")

        text = await cls._extract_text_by_content_type(content, source, original_filename)

        if not text.strip():
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Document appears to be empty or contains no extractable text"
            )

        return text, doc_id, original_filename, sanitized_filename

    @staticmethod
    async def _extract_filename_from_url(url: str) -> Tuple[str, str]:
        """Extract filename from URL"""
        try:
            if 'drive.google.com' in url:
                file_id = None
                patterns = [
                    r'/file/d/([a-zA-Z0-9-_]+)',
                    r'id=([a-zA-Z0-9-_]+)',
                    r'/d/([a-zA-Z0-9-_]+)'
                ]

                for pattern in patterns:
                    match = re.search(pattern, url)
                    if match:
                        file_id = match.group(1)
                        break

                if file_id:
                    metadata = await GoogleDriveProcessor.get_drive_file_metadata(file_id)
                    original_name = metadata["name"]

                    if not original_name.startswith('drive_document_') and not original_name.startswith('gdrive_doc_'):
                        sanitized_name = GoogleDriveProcessor.sanitize_filename(original_name)
                        return original_name, sanitized_name

            parsed = urlparse(url)
            filename = parsed.path.split('/')[-1]
            if filename and '.' in filename:
                sanitized = GoogleDriveProcessor.sanitize_filename(filename)
                return filename, sanitized

        except Exception as e:
            logger.warning(f"Error extracting filename from {url}: {e}")

        timestamp = int(time.time())
        default_name = f"document_{timestamp}.pdf"
        return default_name, default_name

    @classmethod
    async def _extract_text_by_content_type(cls, content: bytes, source: str, filename: str) -> str:
        """Extract text based on content type"""
        if content.startswith(b'%PDF'):
            return cls.extract_text_from_pdf(content)
        elif content.startswith(b'PK'):
            return cls.extract_text_from_docx(content)
        elif b'From:' in content[:1000] or b'Subject:' in content[:1000]:
            return cls.extract_text_from_email(content)

        filename_lower = filename.lower()
        if any(ext in filename_lower for ext in ['.pdf']):
            return cls.extract_text_from_pdf(content)
        elif any(ext in filename_lower for ext in ['.docx', '.doc']):
            return cls.extract_text_from_docx(content)
        elif '.eml' in filename_lower:
            return cls.extract_text_from_email(content)
        else:
            try:
                return content.decode('utf-8', errors='ignore')
            except Exception:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="Unsupported document format"
                )

    @classmethod
    async def process_uploaded_file(cls, content: bytes, filename: str) -> str:
        """Process uploaded file content"""
        filename_lower = filename.lower()

        if filename_lower.endswith('.pdf') or content.startswith(b'%PDF'):
            text = cls.extract_text_from_pdf(content)
        elif filename_lower.endswith('.docx') or content.startswith(b'PK'):
            text = cls.extract_text_from_docx(content)
        elif filename_lower.endswith('.eml'):
            text = cls.extract_text_from_email(content)
        else:
            try:
                text = content.decode('utf-8', errors='ignore')
            except Exception:
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="Unsupported document format"
                )

        if not text.strip():
            raise HTTPException(
                status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                detail="Document appears to be empty"
            )

        return text

    @staticmethod
    def sanitize_filename(filename: str) -> str:
        """Sanitize filename for safe storage"""
        sanitized = filename.replace(" ", "_").replace("/", "_").replace("\\", "_")
        sanitized = re.sub(r'[<>:"|?*]', '_', sanitized)
        sanitized = re.sub(r'_+', '_', sanitized)
        return sanitized.strip('_')


class GoogleDriveProcessor:
    """Handle Google Drive folder and file processing"""

    @staticmethod
    def is_google_drive_link(url: str) -> bool:
        """Check if URL is a Google Drive link"""
        return 'drive.google.com' in url or 'docs.google.com' in url

    @staticmethod
    def extract_folder_id(url: str) -> str:
        """Extract folder ID from Google Drive URL"""
        patterns = [
            r'/folders/([a-zA-Z0-9-_]+)',
            r'id=([a-zA-Z0-9-_]+)',
            r'/d/([a-zA-Z0-9-_]+)'
        ]

        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)

        raise ValueError("Could not extract folder ID from Google Drive URL")

    @staticmethod
    def convert_to_direct_download(url: str) -> str:
        """Convert Google Drive share URL to direct download URL"""
        if '/file/d/' in url:
            file_id_match = re.search(r'/file/d/([a-zA-Z0-9-_]+)', url)
            if file_id_match:
                file_id = file_id_match.group(1)
                return f"https://drive.google.com/uc?export=download&id={file_id}"

        if 'id=' in url:
            parsed = urlparse(url)
            params = parse_qs(parsed.query)
            if 'id' in params:
                file_id = params['id'][0]
                return f"https://drive.google.com/uc?export=download&id={file_id}"

        return url

    @staticmethod
    async def get_drive_file_metadata(file_id: str) -> Dict[str, str]:
        """Get file metadata from Google Drive"""
        try:
            api_url = f"https://www.googleapis.com/drive/v3/files/{file_id}?fields=name,mimeType,size"

            timeout = httpx.Timeout(15.0, connect=5.0)
            async with httpx.AsyncClient(timeout=timeout, follow_redirects=True) as client:
                try:
                    response = await client.get(api_url)
                    if response.status_code == 200:
                        metadata = response.json()
                        original_name = metadata.get("name", "").strip()

                        if original_name and len(original_name) > 0:
                            return {
                                "name": original_name,
                                "mime_type": metadata.get("mimeType", "application/pdf"),
                                "size": metadata.get("size", "unknown")
                            }
                except Exception as api_error:
                    logger.warning(f"Drive API failed: {api_error}")

            fallback_name = f"drive_document_{file_id[:8]}.pdf"
            return {
                "name": fallback_name,
                "mime_type": "application/pdf",
                "size": "unknown"
            }

        except Exception as e:
            logger.error(f"Metadata extraction failed for {file_id}: {e}")
            return {
                "name": f"drive_document_{file_id[:8]}.pdf",
                "mime_type": "application/pdf",
                "size": "unknown"
            }

    @staticmethod
    def sanitize_filename(filename: str) -> str:
        """Sanitize filename for safe storage"""
        sanitized = filename.replace(" ", "_").replace("/", "_").replace("\\", "_")
        sanitized = re.sub(r'[<>:"|?*]', '_', sanitized)
        sanitized = re.sub(r'_+', '_', sanitized)
        return sanitized.strip('_')


class VectorStore:
    """FAISS-based vector store with in-memory caching"""

    def __init__(self):
        self.embedding_model = SentenceTransformer(EMBEDDING_MODEL_NAME)
        self.dimension = self.embedding_model.get_sentence_embedding_dimension()
        self.cache = get_vector_cache()  # Initialize cache
        logger.info(f"✅ Initialized vector store with {EMBEDDING_MODEL_NAME} (dim: {self.dimension})")

    async def create_embeddings(self, document_id: str, chunks: List[Dict[str, Any]],
                                file_name: str = None, source_info: Dict[str, Any] = None) -> int:
        """Create embeddings and store in Supabase + cache"""
        try:
            logger.info(f"Creating embeddings for {len(chunks)} chunks")

            texts = [chunk["text"] for chunk in chunks]

            # Generate embeddings in batches
            batch_size = 32
            all_embeddings = []

            for i in range(0, len(texts), batch_size):
                batch_texts = texts[i:i + batch_size]
                batch_embeddings = self.embedding_model.encode(batch_texts, show_progress_bar=False)
                all_embeddings.append(batch_embeddings)

            embeddings = np.vstack(all_embeddings).astype('float32')

            # Save to Supabase
            await self._save_to_supabase_direct(document_id, embeddings, chunks, file_name, source_info)

            # Cache in memory for instant access
            metadata = {
                "document_id": document_id,
                "chunks_count": len(chunks),
                "embedding_model": EMBEDDING_MODEL_NAME,
                "dimension": self.dimension,
                "created_at": time.time(),
                "total_characters": sum(chunk["char_count"] for chunk in chunks),
                "file_name": file_name,
                "original_filename": source_info.get("original_filename", file_name) if source_info else file_name,
                "sanitized_filename": source_info.get("sanitized_filename", file_name) if source_info else file_name,
                "source_info": source_info or {}
            }
            self.cache.set(document_id, embeddings, chunks, metadata)

            logger.info(f"✅ Created and cached {len(chunks)} vectors")
            return len(chunks)

        except Exception as e:
            logger.error(f"Error creating embeddings: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to create embeddings: {str(e)}"
            )

    async def search_similar_chunks(self, document_id: str, query: str, top_k: int = TOP_K_RETRIEVAL) -> List[
        Dict[str, Any]]:
        """Search for similar chunks using cached data"""
        try:
            # Try to get from cache first
            cached_data = self.cache.get(document_id)

            if cached_data:
                logger.info(f"🚀 Using cached vectors for document {document_id}")
                embeddings = cached_data["embeddings"]
                chunks = cached_data["chunks"]
                index = cached_data["faiss_index"]
            else:
                # Load from Supabase only if not cached
                logger.info(f"📥 Loading vectors from Supabase for document {document_id}")
                embeddings, chunks = await self._load_from_supabase_direct(document_id)

                # Create FAISS index
                index = faiss.IndexFlatIP(self.dimension)
                embeddings_normalized = embeddings.astype('float32')
                faiss.normalize_L2(embeddings_normalized)
                index.add(embeddings_normalized)

                # Cache for future use
                metadata = {"document_id": document_id, "chunks_count": len(chunks)}
                self.cache.set(document_id, embeddings, chunks, metadata)

            # Generate query embedding
            query_embedding = self.embedding_model.encode([query])
            query_embedding = query_embedding.astype('float32')
            faiss.normalize_L2(query_embedding)

            # Search using cached FAISS index
            scores, indices = index.search(query_embedding, min(top_k, index.ntotal))

            # Build results
            retrieved_chunks = []
            for i, (score, idx) in enumerate(zip(scores[0], indices[0])):
                if idx != -1:
                    chunk = chunks[idx].copy()
                    chunk["similarity_score"] = float(score)
                    chunk["rank"] = i + 1
                    retrieved_chunks.append(chunk)

            return retrieved_chunks

        except Exception as e:
            logger.error(f"Error searching chunks: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to search chunks: {str(e)}"
            )

    async def _save_to_supabase_direct(self, document_id: str, embeddings: np.ndarray, chunks: List[Dict[str, Any]],
                                       file_name: str = None, source_info: Dict[str, Any] = None):
        """Save embeddings and chunks to Supabase (unchanged)"""
        try:
            # Save embeddings
            embeddings_bytes = io.BytesIO()
            np.save(embeddings_bytes, embeddings)
            embeddings_data = embeddings_bytes.getvalue()

            embeddings_path = f"vectors/{document_id}/embeddings.npy"
            await upload_file_to_supabase(embeddings_path, embeddings_data)

            # Save chunks
            chunks_json = json.dumps(chunks, indent=2)
            chunks_bytes = chunks_json.encode('utf-8')
            chunks_path = f"vectors/{document_id}/chunks.json"
            await upload_file_to_supabase(chunks_path, chunks_bytes)

            # Save metadata
            metadata = {
                "document_id": document_id,
                "chunks_count": len(chunks),
                "embedding_model": EMBEDDING_MODEL_NAME,
                "dimension": self.dimension,
                "created_at": time.time(),
                "total_characters": sum(chunk["char_count"] for chunk in chunks),
                "file_name": file_name,
                "original_filename": source_info.get("original_filename", file_name) if source_info else file_name,
                "sanitized_filename": source_info.get("sanitized_filename", file_name) if source_info else file_name,
                "source_info": source_info or {}
            }
            metadata_json = json.dumps(metadata, indent=2)
            metadata_bytes = metadata_json.encode('utf-8')
            metadata_path = f"vectors/{document_id}/metadata.json"
            await upload_file_to_supabase(metadata_path, metadata_bytes)

        except Exception as e:
            logger.error(f"Error saving to Supabase: {e}")
            raise

    async def _load_from_supabase_direct(self, document_id: str) -> Tuple[np.ndarray, List[Dict[str, Any]]]:
        """Load embeddings and chunks from Supabase (unchanged)"""
        try:
            # Load embeddings
            embeddings_path = f"vectors/{document_id}/embeddings.npy"
            embeddings_bytes = await download_file_from_supabase(embeddings_path)
            embeddings_io = io.BytesIO(embeddings_bytes)
            embeddings = np.load(embeddings_io)

            # Load chunks
            chunks_path = f"vectors/{document_id}/chunks.json"
            chunks_bytes = await download_file_from_supabase(chunks_path)
            chunks_json = chunks_bytes.decode('utf-8')
            chunks = json.loads(chunks_json)

            return embeddings, chunks

        except Exception as e:
            logger.error(f"Error loading from Supabase: {e}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Vector data not found for document {document_id}"
            )

    async def delete_document_vectors(self, document_id: str) -> bool:
        """Delete document vectors from Supabase AND cache"""
        try:
            # Remove from cache first
            self.cache.remove(document_id)

            # Delete from Supabase
            embeddings_path = f"vectors/{document_id}/embeddings.npy"
            chunks_path = f"vectors/{document_id}/chunks.json"
            metadata_path = f"vectors/{document_id}/metadata.json"

            results = []
            for path in [embeddings_path, chunks_path, metadata_path]:
                try:
                    result = await delete_file_from_supabase(path)
                    results.append(result)
                except Exception as e:
                    logger.warning(f"Failed to delete {path}: {e}")
                    results.append(False)

            return any(results)

        except Exception as e:
            logger.error(f"Error deleting document vectors: {e}")
            return False

    async def list_stored_documents(self) -> List[Dict[str, Any]]:
        """List all documents - use cache if initialized, otherwise load from Supabase"""
        try:
            # If cache is initialized, return cached metadata (instant response)
            if self.cache.is_initialized():
                logger.info("🚀 Returning documents from cache")
                return self.cache.get_all_metadata()

            # Otherwise, load from Supabase and populate cache
            logger.info("📥 Loading documents from Supabase (first time)")
            files = await list_supabase_files(prefix="vectors/")

            documents = []
            seen_doc_ids = set()

            for file_info in files:
                file_path = file_info.get('name', '')

                if file_path.startswith('vectors/') and file_path.endswith('/metadata.json'):
                    try:
                        path_parts = file_path.split('/')
                        if len(path_parts) >= 3:
                            document_id = path_parts[1]

                            if document_id in seen_doc_ids:
                                continue
                            seen_doc_ids.add(document_id)

                            metadata_bytes = await download_file_from_supabase(file_path)
                            metadata = json.loads(metadata_bytes.decode('utf-8'))

                            documents.append({
                                "document_id": document_id,
                                "file_name": metadata.get("file_name", "unknown"),
                                "chunks_count": metadata.get("chunks_count", 0),
                                "total_characters": metadata.get("total_characters", 0),
                                "embedding_model": metadata.get("embedding_model", "unknown"),
                                "created_at": metadata.get("created_at", 0),
                                "status": "stored_in_supabase",
                                "supabase_path": f"vectors/{document_id}/"
                            })

                    except Exception as e:
                        logger.warning(f"Failed to process metadata file {file_path}: {e}")
                        continue

            return documents

        except Exception as e:
            logger.error(f"Error listing documents: {e}")
            return []

    async def preload_all_vectors(self):
        """
        Preload all vectors into cache at startup
        Call this during application startup for best performance
        """
        try:
            logger.info("🔄 Preloading all vectors into cache...")

            files = await list_supabase_files(prefix="vectors/")
            seen_doc_ids = set()

            for file_info in files:
                file_path = file_info.get('name', '')

                if file_path.startswith('vectors/') and file_path.endswith('/metadata.json'):
                    try:
                        path_parts = file_path.split('/')
                        if len(path_parts) >= 3:
                            document_id = path_parts[1]

                            if document_id in seen_doc_ids or self.cache.is_cached(document_id):
                                continue
                            seen_doc_ids.add(document_id)

                            # Load embeddings, chunks, and metadata
                            logger.info(f"Loading document {document_id}...")
                            embeddings, chunks = await self._load_from_supabase_direct(document_id)

                            metadata_bytes = await download_file_from_supabase(file_path)
                            metadata = json.loads(metadata_bytes.decode('utf-8'))

                            # Cache it
                            self.cache.set(document_id, embeddings, chunks, metadata)

                    except Exception as e:
                        logger.warning(f"Failed to preload document {document_id}: {e}")
                        continue

            self.cache.mark_initialized()
            stats = self.cache.get_stats()
            logger.info(f"✅ Preloading complete: {stats['total_documents']} documents, {stats['total_chunks']} chunks")

        except Exception as e:
            logger.error(f"Error preloading vectors: {e}")

    async def search_across_documents(self, query: str, top_k: int = 10, max_docs: int = 5) -> List[Dict[str, Any]]:
        """Search across all documents using cached data"""
        try:
            # Get all documents (will use cache if available)
            all_documents = await self.list_stored_documents()

            if not all_documents:
                return []

            # Shortlist documents
            shortlisted_docs = sorted(
                all_documents,
                key=lambda x: (x.get('created_at', 0), x.get('total_characters', 0)),
                reverse=True
            )[:max_docs]

            # Search tasks
            search_tasks = []
            for doc in shortlisted_docs:
                task = self._search_single_document_with_metadata(doc, query, top_k)
                search_tasks.append(task)

            results = await asyncio.gather(*search_tasks, return_exceptions=True)

            all_chunks = []
            for i, result in enumerate(results):
                if isinstance(result, Exception):
                    logger.warning(f"Search failed for document {shortlisted_docs[i]['document_id']}: {result}")
                    continue
                all_chunks.extend(result)

            all_chunks.sort(key=lambda x: x['similarity_score'], reverse=True)
            return all_chunks[:top_k]

        except Exception as e:
            logger.error(f"Error in global search: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to perform global search: {str(e)}"
            )

    async def search_filtered_documents(self, query: str, document_ids: List[str], top_k: int = 10) -> List[
        Dict[str, Any]]:
        """Search across specific documents using cached data"""
        try:
            all_documents = await self.list_stored_documents()

            if not all_documents:
                return []

            filtered_docs = []
            for doc in all_documents:
                doc_id = doc["document_id"]
                file_name = doc.get("file_name", "unknown")

                if (doc_id in document_ids or
                        file_name in document_ids or
                        any(doc_id.startswith(did) for did in document_ids if len(did) >= 8)):
                    filtered_docs.append(doc)

            if not filtered_docs:
                return []

            search_tasks = []
            for doc in filtered_docs:
                task = self._search_single_document_with_metadata(doc, query, top_k)
                search_tasks.append(task)

            results = await asyncio.gather(*search_tasks, return_exceptions=True)

            all_chunks = []
            for i, result in enumerate(results):
                if isinstance(result, Exception):
                    logger.warning(f"Search failed for document {filtered_docs[i]['document_id']}: {result}")
                    continue
                all_chunks.extend(result)

            all_chunks.sort(key=lambda x: x['similarity_score'], reverse=True)
            return all_chunks[:top_k]

        except Exception as e:
            logger.error(f"Error in filtered document search: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to perform filtered document search: {str(e)}"
            )

    async def _search_single_document_with_metadata(self, doc_metadata: Dict[str, Any], query: str, top_k: int) -> List[
        Dict[str, Any]]:
        """Search a single document using cached data"""
        try:
            document_id = doc_metadata['document_id']

            # Try cache first
            cached_data = self.cache.get(document_id)

            if cached_data:
                embeddings = cached_data["embeddings"]
                chunks = cached_data["chunks"]
                index = cached_data["faiss_index"]
            else:
                # Load from Supabase if not cached
                embeddings, chunks = await self._load_from_supabase_direct(document_id)

                index = faiss.IndexFlatIP(self.dimension)
                embeddings_normalized = embeddings.astype('float32')
                faiss.normalize_L2(embeddings_normalized)
                index.add(embeddings_normalized)

                # Cache it
                metadata = {"document_id": document_id, "chunks_count": len(chunks)}
                self.cache.set(document_id, embeddings, chunks, metadata)

            # Query
            query_embedding = self.embedding_model.encode([query])
            query_embedding = query_embedding.astype('float32')
            faiss.normalize_L2(query_embedding)

            scores, indices = index.search(query_embedding, min(top_k, index.ntotal))

            results = []
            for score, idx in zip(scores[0], indices[0]):
                if idx != -1 and score > 0.3:
                    chunk = chunks[idx].copy()
                    chunk.update({
                        "similarity_score": float(score),
                        "document_id": document_id,
                        "file_name": doc_metadata.get('file_name', 'unknown'),
                        "chunk_preview": chunk["text"][:150] + "..." if len(chunk["text"]) > 150 else chunk["text"]
                    })
                    results.append(chunk)

            return results

        except Exception as e:
            logger.warning(f"Failed to search document {document_id}: {e}")
            return []

    async def get_document_chunks(self, document_id: str) -> List[Dict[str, Any]]:
        """Get chunks for a specific document (use cache if available)"""
        try:
            cached_data = self.cache.get(document_id)
            if cached_data:
                return cached_data["chunks"]

            # Load from Supabase
            chunks_path = f"vectors/{document_id}/chunks.json"
            chunks_bytes = await download_file_from_supabase(chunks_path)
            chunks_json = chunks_bytes.decode('utf-8')
            chunks = json.loads(chunks_json)
            return chunks
        except Exception as e:
            logger.error(f"Error getting chunks: {e}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Chunks not found for document {document_id}"
            )

# API Routes
@app.get("/")
async def list_all_documents():
    """Root endpoint - List all available documents"""
    try:
        documents = await vector_store.list_stored_documents()

        formatted_docs = []
        for doc in documents:
            formatted_docs.append({
                "document_id": doc["document_id"],
                "file_name": doc.get("file_name", "unknown"),
                "chunks_count": doc.get("chunks_count", 0),
                "created_at": int(doc.get("created_at", 0))
            })

        return formatted_docs

    except Exception as e:
        logger.error(f"Error fetching documents: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to fetch documents: {str(e)}"
        )


@app.get("/api/v1/health")
async def health_check():
    """Enhanced health check with cache statistics"""
    supabase_status = "unknown"
    try:
        supabase_manager = get_supabase_manager()
        await supabase_manager.list_files()
        supabase_status = "connected"
    except Exception as e:
        supabase_status = f"error: {str(e)[:100]}"

    llm_info = llm_service.get_service_info() if llm_service else {}
    cache_stats = vector_store.cache.get_stats() if vector_store else {}

    return {
        "status": "healthy",
        "timestamp": time.time(),
        "environment": os.getenv("ENVIRONMENT", "development"),
        "platform": "google-cloud-run",
        "services": {
            "llm_provider": "OpenRouter",
            "llm_available": llm_service.client is not None if llm_service else False,
            "llm_model": llm_info.get("model_name"),
            "llm_model_display": llm_info.get("model_display_name"),
            "embedding_model": EMBEDDING_MODEL_NAME,
            "vector_store_ready": vector_store is not None,
            "supabase_status": supabase_status,
            "cache_initialized": cache_stats.get("initialized", False),
            "cached_documents": cache_stats.get("total_documents", 0),
            "cached_chunks": cache_stats.get("total_chunks", 0)
        },
        "configuration": {
            "chunk_size": CHUNK_SIZE,
            "chunk_overlap": CHUNK_OVERLAP,
            "top_k_retrieval": TOP_K_RETRIEVAL,
            "max_context_length": MAX_CONTEXT_LENGTH,
            "supabase_bucket": os.getenv("SUPABASE_BUCKET", "documents"),
            "storage_mode": "SUPABASE_WITH_MEMORY_CACHE"
        },
        "cache_stats": cache_stats
    }


@app.post("/api/v1/documents/upload")
async def upload_document(
        file: UploadFile = File(...),
        token: str = Depends(verify_token)
):
    """Upload and process document for RAG with automatic caching"""
    try:
        document_id = str(uuid.uuid4())
        timestamp = int(time.time())
        sanitized_filename = DocumentProcessor.sanitize_filename(file.filename)
        supabase_file_path = f"documents/{document_id}_{timestamp}_{sanitized_filename}"

        content = await file.read()

        uploaded_path = await upload_file_to_supabase(supabase_file_path, content)
        public_url = await get_public_url(uploaded_path)

        processor = DocumentProcessor()
        text = await processor.process_uploaded_file(content, file.filename)

        chunks = processor.intelligent_chunking(text)

        source_info = {
            "source_type": "file_upload",
            "upload_timestamp": timestamp,
            "original_content_type": file.content_type,
            "original_filename": file.filename,
            "sanitized_filename": sanitized_filename,
            "supabase_path": uploaded_path,
            "file_size_bytes": len(content),
            "ingested_at": time.time(),
            "upload_method": "multipart_form"
        }

        # This will automatically cache the document
        chunks_created = await vector_store.create_embeddings(document_id, chunks, file.filename, source_info)

        logger.info(f"✅ Document uploaded and cached: {document_id}")

        return DocumentUploadResponse(
            document_id=document_id,
            filename=file.filename,
            status="processed_and_cached",
            chunks_created=chunks_created,
            message=f"Document processed and cached successfully with {chunks_created} chunks",
            supabase_path=uploaded_path,
            public_url=public_url if public_url else None
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document upload failed: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to process document: {str(e)}"
        )


@app.post("/api/v1/hackrx/run")
async def process_document_qa_rag(
        request: DocumentQARequest,
        token: str = Depends(verify_token)
):
    """RAG-powered document QA endpoint with OpenRouter and caching"""
    try:
        document_id = None

        # Handle document processing
        if request.documents and not request.document_id:
            text, document_id, original_filename, sanitized_filename = await DocumentProcessor.process_document_from_source(
                request.documents)

            # Check if already cached first (fastest)
            if vector_store.cache.is_cached(document_id):
                logger.info(f"✅ Document already in cache: {document_id}")
            else:
                # Check Supabase
                try:
                    await vector_store._load_from_supabase_direct(document_id)
                    logger.info(f"✅ Using existing vectors from Supabase for: {document_id}")
                    # Load into cache for future use
                    embeddings, chunks = await vector_store._load_from_supabase_direct(document_id)
                    metadata = {"document_id": document_id, "file_name": original_filename}
                    vector_store.cache.set(document_id, embeddings, chunks, metadata)
                except HTTPException:
                    # Create new vectors
                    logger.info(f"📝 Creating new vectors for document: {document_id}")
                    chunks = DocumentProcessor.intelligent_chunking(text)
                    source_info = {
                        "source_url": request.documents,
                        "source_type": "api_request",
                        "processed_via": "hackrx_endpoint",
                        "original_filename": original_filename,
                        "sanitized_filename": sanitized_filename,
                        "ingested_at": time.time()
                    }
                    # This will automatically cache
                    await vector_store.create_embeddings(document_id, chunks, original_filename, source_info)

        elif request.document_id:
            document_id = request.document_id
            logger.info(f"📄 Using provided document_id: {document_id}")
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Either 'documents' URL or 'document_id' must be provided"
            )

        # Helper function to process a single question
        async def process_single_question(question: str, index: int) -> Dict[str, Any]:
            """Process a single question and return answer with sources"""
            logger.info(f"🔍 Processing question {index + 1}/{len(request.questions)}: {question[:50]}...")

            # This will use cached data automatically
            relevant_chunks = await vector_store.search_similar_chunks(
                document_id, question, TOP_K_RETRIEVAL
            )

            # Generate answer using OpenRouter
            answer_data = await llm_service.generate_rag_answer(
                question, relevant_chunks, document_id
            )

            # Build sources list
            sources = []
            for chunk in relevant_chunks[:5]:
                sources.append({
                    "document_id": document_id,
                    "chunk_id": chunk["chunk_id"],
                    "similarity_score": chunk["similarity_score"],
                    "chunk_preview": chunk["text"][:150] + "..." if len(chunk["text"]) > 150 else chunk["text"],
                    "rank": chunk.get("rank", len(sources) + 1)
                })

            # Build detailed answer
            detailed_answer = {
                "question": question,
                "answer": answer_data.get("answer", "No answer generated"),
                "confidence": answer_data.get("confidence", 0.0),
                "sources": sources,
                "chunks_retrieved": len(relevant_chunks),
                "model_used": answer_data.get("model_used"),
                "retrieval_info": {
                    "top_similarity_score": relevant_chunks[0]["similarity_score"] if relevant_chunks else 0.0,
                    "avg_similarity_score": sum(chunk["similarity_score"] for chunk in relevant_chunks) / len(
                        relevant_chunks) if relevant_chunks else 0.0
                }
            }

            logger.info(f"✅ Completed question {index + 1}")
            return detailed_answer

        # Process all questions concurrently for maximum speed
        logger.info(f"🚀 Processing {len(request.questions)} questions in parallel (using cache)...")
        detailed_answers = await asyncio.gather(*[
            process_single_question(question, i)
            for i, question in enumerate(request.questions)
        ])

        # Build final response
        response = {
            "answers": detailed_answers,
            "document_id": document_id,
            "total_questions": len(request.questions),
            "search_type": "document_specific",
            "cache_hit": vector_store.cache.is_cached(document_id),
            "response_format": "markdown",  # Add this line
            "processing_info": {
                "llm_service": "OpenRouter",
                "model_used": llm_service.model_name,
                "total_chunks_searched": sum(len(ans["sources"]) for ans in detailed_answers),
                "avg_confidence": sum(ans["confidence"] for ans in detailed_answers) / len(
                    detailed_answers) if detailed_answers else 0.0
            }
        }

        logger.info(f"✅ Successfully processed all {len(request.questions)} questions")
        return response

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"❌ Error in RAG document QA: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Internal server error: {str(e)}"
        )


@app.post("/api/v1/query/global")
async def global_query(
        request: GlobalQueryRequest,
        token: str = Depends(verify_token)
):
    """Global search across documents with OpenRouter"""
    try:
        if request.document_ids:
            relevant_chunks = await vector_store.search_filtered_documents(
                request.query, request.document_ids, top_k=request.top_k
            )
            search_type = "filtered"
        else:
            relevant_chunks = await vector_store.search_across_documents(
                request.query, top_k=request.top_k, max_docs=request.max_docs
            )
            search_type = "global"

        if not relevant_chunks:
            return {
                "answer": "No relevant information found.",
                "sources": [],
                "search_type": search_type
            }

        answer_data = await llm_service.generate_rag_answer(
            request.query,
            relevant_chunks,
            "global_search"
        )

        sources = []
        for chunk in relevant_chunks[:5]:
            sources.append({
                "document_id": chunk["document_id"],
                "file_name": chunk.get("file_name", "unknown"),
                "chunk_preview": chunk.get("chunk_preview", chunk["text"][:150] + "..."),
                "similarity_score": chunk["similarity_score"]
            })

        unique_doc_count = len(set(chunk["document_id"] for chunk in relevant_chunks))

        return {
            "answer": answer_data.get("answer", "No answer generated"),
            "sources": sources,
            "query": request.query,
            "search_type": search_type,
            "chunks_searched": len(relevant_chunks),
            "documents_searched": unique_doc_count,
            "model_used": answer_data.get("model_used"),
            "filter_applied": request.document_ids is not None
        }

    except Exception as e:
        logger.error(f"Error in global query: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to process global query: {str(e)}"
        )


@app.get("/api/v1/test-llm")
async def test_llm():
    """Test OpenRouter LLM connectivity"""
    try:
        if not llm_service:
            return {"status": "error", "message": "OpenRouter service not initialized"}

        result = await llm_service.test_connection()
        return result

    except Exception as e:
        return {
            "status": "error",
            "message": f"OpenRouter test failed: {str(e)}"
        }


@app.get("/api/v1/models")
async def list_models():
    """List available OpenRouter models"""
    return {
        "current_model": llm_service.model_name if llm_service else None,
        "available_models": AVAILABLE_MODELS
    }


@app.post("/api/v1/models/switch")
async def switch_model(
        model_id: str,
        token: str = Depends(verify_token)
):
    """Hot-switch the active OpenRouter model without restarting"""
    try:
        if not llm_service:
            raise HTTPException(status_code=503, detail="LLM service not initialized")

        old_model = llm_service.model_name
        success = llm_service.switch_model(model_id)

        if success:
            return {
                "status": "success",
                "old_model": old_model,
                "new_model": model_id,
                "display_name": AVAILABLE_MODELS.get(model_id, model_id)
            }
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Failed to switch to model: {model_id}. Reverted to {old_model}."
            )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Model switch failed: {str(e)}"
        )


@app.delete("/api/v1/documents/{document_id}")
async def delete_document(
        document_id: str,
        token: str = Depends(verify_token)
):
    """Delete a document from Supabase and cache"""
    try:
        # This will remove from both cache and Supabase
        success = await vector_store.delete_document_vectors(document_id)

        return {
            "message": f"Document {document_id} deleted from cache and Supabase",
            "document_id": document_id,
            "success": success
        }

    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to delete document: {str(e)}"
        )


@app.get("/api/v1/documents")
async def list_documents(token: str = Depends(verify_token)):
    """List all processed documents"""
    try:
        documents = await vector_store.list_stored_documents()

        return {
            "total_documents": len(documents),
            "storage_mode": "supabase_only",
            "documents": documents
        }

    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to list documents: {str(e)}"
        )


@app.middleware("http")
async def cloud_run_request_middleware(request: Request, call_next):
    """Enhanced request logging middleware"""
    start_time = time.time()

    if killer.kill_now:
        return JSONResponse(
            status_code=503,
            content={"error": "Service is shutting down"}
        )

    path = request.url.path
    method = request.method
    logger.info(f"Request: {method} {path}")

    try:
        response = await call_next(request)
        process_time = time.time() - start_time
        logger.info(f"Response: {response.status_code} ({process_time:.3f}s)")

        response.headers["X-Cloud-Run-Service"] = "rag-document-qa-api"
        response.headers["X-Response-Time"] = f"{process_time:.3f}s"

        return response

    except Exception as e:
        process_time = time.time() - start_time
        logger.error(f"Request failed: {method} {path} - {str(e)} ({process_time:.3f}s)")
        raise


@app.post("/api/v1/cache/refresh")
async def refresh_cache(token: str = Depends(verify_token)):
    """
    Manually refresh the vector cache
    Useful after uploading new documents
    """
    try:
        logger.info("🔄 Manual cache refresh requested")

        # Clear existing cache
        vector_store.cache.clear()

        # Preload all vectors
        await vector_store.preload_all_vectors()

        # Get stats
        stats = vector_store.cache.get_stats()

        return {
            "message": "Cache refreshed successfully",
            "stats": stats
        }
    except Exception as e:
        logger.error(f"Error refreshing cache: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to refresh cache: {str(e)}"
        )


@app.get("/api/v1/cache/stats")
async def get_cache_stats(token: str = Depends(verify_token)):
    """Get current cache statistics"""
    try:
        stats = vector_store.cache.get_stats()
        return {
            "cache_stats": stats,
            "embedding_model": EMBEDDING_MODEL_NAME
        }
    except Exception as e:
        logger.error(f"Error getting cache stats: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get cache stats: {str(e)}"
        )


@app.post("/api/v1/cache/clear")
async def clear_cache(token: str = Depends(verify_token)):
    """Clear the entire cache (memory only, Supabase data remains)"""
    try:
        vector_store.cache.clear()
        return {
            "message": "Cache cleared successfully",
            "note": "Supabase data remains intact"
        }
    except Exception as e:
        logger.error(f"Error clearing cache: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to clear cache: {str(e)}"
        )

@app.exception_handler(HTTPException)
async def http_exception_handler(request: Request, exc: HTTPException):
    """HTTP exception handler"""
    logger.error(f"HTTP Exception [{request.url.path}]: {exc.status_code} - {exc.detail}")
    return JSONResponse(
        status_code=exc.status_code,
        content={
            "error": exc.detail,
            "status_code": exc.status_code,
            "path": request.url.path
        }
    )


@app.exception_handler(Exception)
async def general_exception_handler(request: Request, exc: Exception):
    """General exception handler"""
    logger.error(f"Unhandled exception [{request.url.path}]: {exc}", exc_info=True)
    return JSONResponse(
        status_code=500,
        content={
            "error": "Internal server error",
            "details": str(exc) if os.getenv("DEBUG") else None,
            "path": request.url.path
        }
    )


# ============================================================================
# TEMPORARY CHAT SESSION ENDPOINTS
# ============================================================================

@app.post("/api/v1/temporary/session/create")
async def create_temporary_session(
        request: TemporarySessionCreate,
        token: str = Depends(verify_token)
):
    """Create a new temporary chat session"""
    try:
        session_manager = get_session_manager()
        session_id = session_manager.create_session(ttl_minutes=request.ttl_minutes)

        session = session_manager.get_session(session_id)

        return {
            "session_id": session_id,
            "status": "created",
            "ttl_minutes": request.ttl_minutes,
            "expires_at": session.expires_at,
            "message": f"Temporary session created with {request.ttl_minutes} minute TTL"
        }

    except Exception as e:
        logger.error(f"Error creating temporary session: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to create session: {str(e)}"
        )


@app.post("/api/v1/temporary/upload")
async def upload_temporary_document(
        session_id: str,
        file: UploadFile = File(...),
        token: str = Depends(verify_token)
):
    """Upload document to temporary session (in-memory only, no DB/Supabase)"""
    try:
        session_manager = get_session_manager()
        session = session_manager.get_session(session_id)

        if not session:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Session not found or expired"
            )

        # Generate document ID
        document_id = str(uuid.uuid4())

        # Read file content
        content = await file.read()

        # Process document
        processor = DocumentProcessor()
        text = await processor.process_uploaded_file(content, file.filename)

        # Create chunks
        chunks = processor.intelligent_chunking(text)

        # Create embeddings (in-memory only)
        texts = [chunk["text"] for chunk in chunks]
        embeddings = vector_store.embedding_model.encode(texts, show_progress_bar=False)

        # Store in session (no Supabase, no persistent storage)
        metadata = {
            "original_filename": file.filename,
            "file_size": len(content),
            "content_type": file.content_type,
            "session_id": session_id
        }

        session.add_document(
            doc_id=document_id,
            filename=file.filename,
            chunks=chunks,
            embeddings=embeddings,
            metadata=metadata
        )

        logger.info(f"✅ Uploaded temporary document {document_id} to session {session_id}")

        return {
            "document_id": document_id,
            "filename": file.filename,
            "session_id": session_id,
            "status": "uploaded_to_session",
            "chunks_created": len(chunks),
            "storage_type": "temporary_in_memory",
            "message": f"Document uploaded to temporary session (will be deleted when session expires)",
            "session_info": session.get_stats()
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error uploading temporary document: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to upload document: {str(e)}"
        )


@app.post("/api/v1/temporary/query")
async def query_temporary_documents(
        request: TemporaryDocumentQARequest,
        token: str = Depends(verify_token)
):
    """Query documents in temporary session (uses in-memory data only)"""
    try:
        session_manager = get_session_manager()
        session = session_manager.get_session(request.session_id)

        if not session:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Session not found or expired"
            )

        # Get document to query
        if request.document_id:
            doc_data = session.get_document(request.document_id)
            if not doc_data:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail=f"Document {request.document_id} not found in session"
                )
            documents_to_query = [doc_data]
        else:
            # Query all documents in session
            doc_list = session.list_documents()
            if not doc_list:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="No documents in session"
                )
            documents_to_query = [session.get_document(doc["document_id"]) for doc in doc_list]

        # Process questions
        async def process_question(question: str, index: int) -> Dict[str, Any]:
            """Process single question across session documents"""
            all_results = []

            # Search across all documents in session
            for doc_data in documents_to_query:
                chunks = doc_data["chunks"]
                embeddings = doc_data["embeddings"]
                doc_id = doc_data["metadata"]["document_id"]

                # Create temporary FAISS index
                import faiss
                import numpy as np

                dimension = embeddings.shape[1]
                index = faiss.IndexFlatIP(dimension)
                embeddings_normalized = embeddings.astype('float32')
                faiss.normalize_L2(embeddings_normalized)
                index.add(embeddings_normalized)

                # Query
                query_embedding = vector_store.embedding_model.encode([question])
                query_embedding = query_embedding.astype('float32')
                faiss.normalize_L2(query_embedding)

                scores, indices = index.search(query_embedding, min(TOP_K_RETRIEVAL, len(chunks)))

                # Collect results
                for score, idx in zip(scores[0], indices[0]):
                    if idx != -1:
                        chunk = chunks[idx].copy()
                        chunk["similarity_score"] = float(score)
                        chunk["document_id"] = doc_id
                        chunk["filename"] = doc_data["metadata"]["filename"]
                        all_results.append(chunk)

            # Sort by score
            all_results.sort(key=lambda x: x["similarity_score"], reverse=True)
            relevant_chunks = all_results[:TOP_K_RETRIEVAL]

            # Generate answer
            answer_data = await llm_service.generate_rag_answer(
                question, relevant_chunks, request.session_id
            )

            return {
                "question": question,
                "answer": answer_data.get("answer", "No answer generated"),
                "confidence": answer_data.get("confidence", 0.0),
                "sources": [
                    {
                        "document_id": chunk["document_id"],
                        "filename": chunk.get("filename", "unknown"),
                        "chunk_id": chunk["chunk_id"],
                        "similarity_score": chunk["similarity_score"],
                        "chunk_preview": chunk["text"][:150] + "..."
                    }
                    for chunk in relevant_chunks[:5]
                ],
                "chunks_retrieved": len(relevant_chunks)
            }

        # Process all questions
        answers = await asyncio.gather(*[
            process_question(q, i) for i, q in enumerate(request.questions)
        ])

        return {
            "answers": answers,
            "session_id": request.session_id,
            "storage_type": "temporary_in_memory",
            "total_questions": len(request.questions),
            "session_info": session.get_stats()
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error querying temporary documents: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to query documents: {str(e)}"
        )


@app.get("/api/v1/temporary/session/{session_id}")
async def get_temporary_session(
        session_id: str,
        token: str = Depends(verify_token)
):
    """Get information about a temporary session"""
    try:
        session_manager = get_session_manager()
        session = session_manager.get_session(session_id)

        if not session:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Session not found or expired"
            )

        return {
            "session": session.get_stats(),
            "documents": session.list_documents()
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting session: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get session: {str(e)}"
        )


@app.delete("/api/v1/temporary/session/{session_id}")
async def delete_temporary_session(
        session_id: str,
        token: str = Depends(verify_token)
):
    """Manually delete a temporary session"""
    try:
        session_manager = get_session_manager()
        success = session_manager.delete_session(session_id)

        if not success:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Session not found"
            )

        return {
            "message": f"Session {session_id} deleted successfully",
            "session_id": session_id
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting session: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to delete session: {str(e)}"
        )


@app.post("/api/v1/temporary/session/{session_id}/extend")
async def extend_temporary_session(
        session_id: str,
        minutes: int = 30,
        token: str = Depends(verify_token)
):
    """Extend session TTL"""
    try:
        session_manager = get_session_manager()
        success = session_manager.extend_session(session_id, minutes)

        if not success:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Session not found or expired"
            )

        session = session_manager.get_session(session_id)

        return {
            "message": f"Session extended by {minutes} minutes",
            "session_id": session_id,
            "new_expiry": session.expires_at,
            "time_remaining_seconds": int(session.expires_at - time.time())
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error extending session: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to extend session: {str(e)}"
        )


@app.get("/api/v1/temporary/sessions/stats")
async def get_all_sessions_stats(token: str = Depends(verify_token)):
    """Get statistics for all temporary sessions"""
    try:
        session_manager = get_session_manager()
        return session_manager.get_all_sessions_stats()

    except Exception as e:
        logger.error(f"Error getting sessions stats: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get sessions stats: {str(e)}"
        )

@app.on_event("startup")
async def startup_event():
    """Initialize services and preload vectors on startup"""
    global vector_store, llm_service

    logger.info("🚀 Starting RAG-Powered Document QA API with OpenRouter")
    logger.info("📡 NO LOCAL STORAGE - All data stored in Supabase")
    logger.info(f"🧠 Embedding model: {EMBEDDING_MODEL_NAME}")
    logger.info(f"⚙️ Chunk size: {CHUNK_SIZE}, Overlap: {CHUNK_OVERLAP}")
    logger.info(f"🔍 Top-K retrieval: {TOP_K_RETRIEVAL}")
    logger.info(f"☁️ Supabase bucket: {os.getenv('SUPABASE_BUCKET', 'documents')}")

    # Initialize VectorStore
    try:
        vector_store = VectorStore()
        logger.info("✅ Vector store initialized")
    except Exception as e:
        logger.error(f"❌ Failed to initialize vector store: {e}")
        raise

    # Initialize OpenRouter LLM Service
    try:
        llm_service = OpenRouterService(
            api_key=os.getenv("OPENROUTER_API_KEY"),
            model_name=os.getenv("OPENROUTER_MODEL"),
            app_name=os.getenv("OPENROUTER_APP_NAME", "Microsoft RAG QA"),
        )
        logger.info(f"✅ OpenRouter service initialized with model: {llm_service.model_name}")
    except Exception as e:
        logger.error(f"❌ Failed to initialize OpenRouter service: {e}")
        raise

    # Initialize Session Manager for temporary chat
    try:
        session_manager = get_session_manager()
        await session_manager.start_cleanup_task()
        logger.info("✅ Session manager initialized with automatic cleanup")
    except Exception as e:
        logger.warning(f"⚠️ Failed to initialize session manager: {e}")

    # Preload all vectors into memory cache
    try:
        logger.info("🔄 Preloading vectors into memory cache...")
        await vector_store.preload_all_vectors()
        stats = vector_store.cache.get_stats()
        logger.info(f"✅ Cache preloaded: {stats['total_documents']} documents, {stats['total_chunks']} chunks")
    except Exception as e:
        logger.warning(f"⚠️ Failed to preload vectors (will load on-demand): {e}")

    logger.info("🚀 Startup initialization complete")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
